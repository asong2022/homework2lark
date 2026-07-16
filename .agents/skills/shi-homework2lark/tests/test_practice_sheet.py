from __future__ import annotations

import importlib.util
import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

SKILL_ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = SKILL_ROOT / "scripts" / "practice_sheet.py"
SPEC = importlib.util.spec_from_file_location("practice_sheet", MODULE_PATH)
if SPEC is None or SPEC.loader is None:
    raise RuntimeError("Unable to load practice_sheet module")
practice_sheet = importlib.util.module_from_spec(SPEC)
sys.modules[SPEC.name] = practice_sheet
SPEC.loader.exec_module(practice_sheet)

WPS_TEXTBOX = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"


def _manifest() -> dict[str, object]:
    return {
        "batchCode": "20260715-01",
        "manifestVersion": "practice-sheet-v1",
        "pages": [
            {
                "pageNumber": 1,
                "items": [
                    {
                        "itemCode": "R01",
                        "question": "在括号里填合适的单位。",
                        "answerLines": 0,
                        "source": {
                            "type": "original",
                            "questionId": "problem_units",
                        },
                    }
                ],
            },
            {
                "pageNumber": 2,
                "items": [
                    {
                        "itemCode": "R02",
                        "question": "在数轴下面的方框里填上合适的小数。",
                        "stemImage": {
                            "path": "images/number-line.png",
                            "description": "0 到 2 的数轴和三个待填写方框",
                        },
                        "answerLines": 1,
                        "source": {
                            "type": "variant",
                            "questionId": "problem_number_line",
                            "variantId": "variant:v1:abc",
                        },
                    }
                ],
            },
        ],
    }


def _personal_manifest() -> dict[str, object]:
    manifest = _manifest()
    manifest["manifestVersion"] = "personal-practice-v1"
    manifest["student"] = {
        "name": "学生A",
        "studentNumber": "44",
        "instanceCode": "S001",
    }
    return manifest


def _auto_manifest() -> dict[str, object]:
    legacy = _manifest()
    return {
        "batchCode": legacy["batchCode"],
        "manifestVersion": "practice-sheet-v2",
        "items": [item for page in legacy["pages"] for item in page["items"]],
    }


def _auto_personal_manifest() -> dict[str, object]:
    manifest = _auto_manifest()
    manifest["manifestVersion"] = "personal-practice-v2"
    manifest["student"] = {
        "name": "学生A",
        "studentNumber": "44",
        "instanceCode": "S001",
    }
    return manifest


class PracticeSheetTests(unittest.TestCase):
    def _workspace(self, manifest=None):
        temporary = tempfile.TemporaryDirectory()
        root = Path(temporary.name)
        image_dir = root / "images"
        image_dir.mkdir()
        Image.new("RGB", (800, 220), "white").save(image_dir / "number-line.png")
        path = root / "manifest.json"
        path.write_text(json.dumps(manifest or _manifest(), ensure_ascii=False), encoding="utf-8")
        return temporary, root, path

    def test_manifest_derives_date_page_codes_images_and_sources(self) -> None:
        temporary, root, path = self._workspace()
        self.addCleanup(temporary.cleanup)

        manifest = practice_sheet._load_manifest(path)

        self.assertEqual(manifest.batch_code, "20260715-01")
        self.assertEqual(practice_sheet._format_title(manifest.batch_date), "7月15日练习纸")
        self.assertEqual([page.page_number for page in manifest.pages], [1, 2])
        self.assertEqual([page.items[0].display_number for page in manifest.pages], [1, 2])
        self.assertEqual(manifest.pages[0].page_code, "20260715-01-P1")
        self.assertEqual(manifest.pages[1].page_code, "20260715-01-P2")
        self.assertEqual(manifest.pages[0].items[0].answer_lines, 0)
        self.assertEqual(
            manifest.pages[1].items[0].stem_image.path,
            (root / "images" / "number-line.png").resolve(),
        )
        self.assertEqual(manifest.pages[1].items[0].source.variant_id, "variant:v1:abc")

    def test_legacy_title_and_qr_fields_are_rejected(self) -> None:
        raw = _manifest()
        raw["title"] = "旧标题"
        raw["pages"][0]["qrPayload"] = "20260715-01-P1"
        temporary, _, path = self._workspace(raw)
        self.addCleanup(temporary.cleanup)

        with self.assertRaisesRegex(practice_sheet.ManifestError, "manifest 字段"):
            practice_sheet._load_manifest(path)

    def test_batch_code_uses_a_valid_date_and_daily_sequence(self) -> None:
        raw = _manifest()
        raw["batchCode"] = "20260230-01"
        temporary, _, path = self._workspace(raw)
        self.addCleanup(temporary.cleanup)

        with self.assertRaisesRegex(practice_sheet.ManifestError, "日期无效"):
            practice_sheet._load_manifest(path)

    def test_item_codes_are_globally_consecutive(self) -> None:
        raw = _manifest()
        raw["pages"][1]["items"][0]["itemCode"] = "R03"
        temporary, _, path = self._workspace(raw)
        self.addCleanup(temporary.cleanup)

        with self.assertRaisesRegex(practice_sheet.ManifestError, "R01 连续"):
            practice_sheet._load_manifest(path)

    def test_variant_source_requires_variant_id(self) -> None:
        raw = _manifest()
        raw["pages"][1]["items"][0]["source"].pop("variantId")
        temporary, _, path = self._workspace(raw)
        self.addCleanup(temporary.cleanup)

        with self.assertRaisesRegex(practice_sheet.ManifestError, "original 或 variant"):
            practice_sheet._load_manifest(path)

    def test_personal_manifest_derives_private_title_and_anonymous_page_code(self) -> None:
        temporary, _, path = self._workspace(_auto_personal_manifest())
        self.addCleanup(temporary.cleanup)

        manifest = practice_sheet._load_manifest(path)

        self.assertIsNotNone(manifest.student)
        self.assertEqual(manifest.student.name, "学生A")
        self.assertEqual(manifest.student.student_number, "44")
        self.assertTrue(manifest.natural_flow)
        self.assertEqual(manifest.page_code_prefix, "20260715-01-S001-P")
        self.assertEqual([item.display_number for item in manifest.items], [1, 2])
        self.assertEqual(
            practice_sheet._format_title(manifest.batch_date, personal=True),
            "7月15日个人练习纸",
        )
        self.assertNotIn("学生A", manifest.page_code_prefix)
        self.assertNotIn("44", manifest.page_code_prefix)

    def test_personal_manifest_rejects_invalid_instance_code(self) -> None:
        raw = _auto_personal_manifest()
        raw["student"]["instanceCode"] = "00"
        temporary, _, path = self._workspace(raw)
        self.addCleanup(temporary.cleanup)

        with self.assertRaisesRegex(practice_sheet.ManifestError, "S 加三位数字"):
            practice_sheet._load_manifest(path)

    def test_standard_manifest_rejects_personal_identity_fields(self) -> None:
        raw = _auto_manifest()
        raw["student"] = _auto_personal_manifest()["student"]
        temporary, _, path = self._workspace(raw)
        self.addCleanup(temporary.cleanup)

        with self.assertRaisesRegex(practice_sheet.ManifestError, "manifest 字段"):
            practice_sheet._load_manifest(path)

    def test_stem_image_must_be_readable_and_supported(self) -> None:
        raw = _manifest()
        raw["pages"][1]["items"][0]["stemImage"]["path"] = "images/missing.bmp"
        temporary, _, path = self._workspace(raw)
        self.addCleanup(temporary.cleanup)

        with self.assertRaisesRegex(practice_sheet.ManifestError, "题干图片"):
            practice_sheet._load_manifest(path)

    def test_builds_formal_two_page_docx_with_one_section_and_hard_page_break(self) -> None:
        temporary, root, path = self._workspace()
        self.addCleanup(temporary.cleanup)
        output = root / "正式练习纸.docx"

        practice_sheet.build_document(path, output)

        self.assertTrue(output.is_file())
        document = Document(output)
        body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("7月15日练习纸", body_text)
        self.assertIn("1.  在括号里填合适的单位。", body_text)
        self.assertIn("2.  在数轴下面的方框里填上合适的小数。", body_text)
        self.assertNotIn("R01", body_text)

        title = next(p for p in document.paragraphs if "7月15日练习纸" in p.text)
        title_run = next(run for run in title.runs if "练习纸" in run.text)
        self.assertEqual(title_run.font.name, "黑体")
        self.assertEqual(title_run.font.size.pt, 16)

        self.assertEqual(len(document.sections), 1)
        self.assertEqual(len(document._element.xpath(".//w:br[@w:type='page']")), 1)
        self.assertEqual(len(document._element.xpath(".//w:sectPr")), 1)

        textboxes = list(document._element.iter(WPS_TEXTBOX))
        textbox_text = [
            "".join(text.text or "" for text in node.iter(qn("w:t"))) for node in textboxes
        ]
        self.assertEqual(textbox_text, ["20260715-01-P1", "20260715-01-P2"])
        anchors = list(document._element.iter(qn("wp:anchor")))
        self.assertEqual(len(anchors), 2)
        for anchor in anchors:
            self.assertEqual(anchor.find(qn("wp:positionH")).get("relativeFrom"), "page")
            self.assertEqual(anchor.find(qn("wp:positionV")).get("relativeFrom"), "page")

        footer = document.sections[0].footer
        footer_text = "\n".join(paragraph.text for paragraph in footer.paragraphs)
        footer_fields = footer._element.xpath(".//w:instrText")
        self.assertNotIn("20260715-01", footer_text)
        self.assertEqual([field.text.strip() for field in footer_fields], ["PAGE"])

        image_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph._p.xpath(".//w:drawing/wp:inline")
        ]
        self.assertEqual(len(image_paragraphs), 1)
        self.assertEqual(image_paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(len(document.inline_shapes), 1)
        self.assertEqual(
            document.inline_shapes[0]._inline.docPr.get("descr"),
            "0 到 2 的数轴和三个待填写方框",
        )

        question_two = next(
            paragraph for paragraph in document.paragraphs if "数轴下面" in paragraph.text
        )
        self.assertTrue(question_two._p.xpath("./w:pPr/w:keepNext"))
        self.assertTrue(image_paragraphs[0]._p.xpath("./w:pPr/w:keepNext"))

    def test_builds_personal_docx_with_prefilled_identity_and_word_auto_flow(self) -> None:
        temporary, root, path = self._workspace(_auto_personal_manifest())
        self.addCleanup(temporary.cleanup)
        output = root / "个人练习纸.docx"

        practice_sheet.build_document(path, output)

        document = Document(output)
        body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("7月15日个人练习纸", body_text)
        self.assertIn("姓名：学生A", body_text)
        self.assertIn("学号：44", body_text)
        self.assertNotIn("姓名：________________", body_text)
        self.assertEqual(len(document._element.xpath(".//w:br[@w:type='page']")), 0)

        self.assertEqual(list(document._element.iter(WPS_TEXTBOX)), [])
        header = document.sections[0].header
        textboxes = list(header._element.iter(WPS_TEXTBOX))
        textbox_text = [
            "".join(text.text or "" for text in node.iter(qn("w:t"))) for node in textboxes
        ]
        self.assertEqual(textbox_text, ["20260715-01-S001-P1"])
        page_fields = header._element.xpath(".//w:fldSimple")
        self.assertEqual(len(page_fields), 1)
        self.assertEqual(page_fields[0].get(qn("w:instr")).strip(), "PAGE")


if __name__ == "__main__":
    unittest.main()
