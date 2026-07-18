from __future__ import annotations

import argparse
import json
import re
import sys
from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

TITLE_FONT = "黑体"
BODY_FONT = "宋体"
CODE_FONT = "Arial"
BLACK = "000000"
GRAY = "555555"
MAX_IMAGE_WIDTH = Mm(145)
MAX_IMAGE_HEIGHT = Mm(65)
SUPPORTED_IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png"}
LEGACY_STANDARD_MANIFEST_VERSION = "practice-sheet-v1"
LEGACY_PERSONAL_MANIFEST_VERSION = "personal-practice-v1"
STANDARD_MANIFEST_VERSION = "practice-sheet-v2"
PERSONAL_MANIFEST_VERSION = "personal-practice-v2"


class ManifestError(ValueError):
    pass


@dataclass(frozen=True)
class PracticeSource:
    source_type: str
    question_id: str
    variant_id: str | None


@dataclass(frozen=True)
class StemImage:
    path: Path
    description: str


@dataclass(frozen=True)
class PracticeItem:
    item_code: str
    display_number: int
    question: str
    answer_lines: int
    source: PracticeSource
    stem_image: StemImage | None


@dataclass(frozen=True)
class PracticePage:
    page_number: int
    page_code: str
    items: tuple[PracticeItem, ...]


@dataclass(frozen=True)
class PracticeStudent:
    name: str
    student_number: str
    instance_code: str


@dataclass(frozen=True)
class PracticeManifest:
    batch_code: str
    batch_date: datetime
    version: str
    student: PracticeStudent | None
    items: tuple[PracticeItem, ...]
    pages: tuple[PracticePage, ...]
    natural_flow: bool
    page_code_prefix: str


def _set_run_font(
    run,
    *,
    font_name: str,
    size: float,
    bold: bool = False,
    color: str = BLACK,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)
    fonts.set(qn("w:cs"), font_name)


def _set_style_font(style, *, font_name: str, size: float, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(BLACK)
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)
    fonts.set(qn("w:cs"), font_name)


def _paragraph_style(document: Document, name: str):
    styles = document.styles
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, font_name=BODY_FONT, size=12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.3

    title = _paragraph_style(document, "Practice Title")
    _set_style_font(title, font_name=TITLE_FONT, size=16, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_together = True

    student = _paragraph_style(document, "Practice Student Info")
    _set_style_font(student, font_name=BODY_FONT, size=12)
    student.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    student.paragraph_format.space_before = Pt(1)
    student.paragraph_format.space_after = Pt(6)
    student.paragraph_format.keep_together = True

    question = _paragraph_style(document, "Practice Question")
    _set_style_font(question, font_name=BODY_FONT, size=12)
    question.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    question.paragraph_format.space_before = Pt(4)
    question.paragraph_format.space_after = Pt(3)
    question.paragraph_format.line_spacing = 1.35
    question.paragraph_format.keep_together = True

    image = _paragraph_style(document, "Practice Stem Image")
    _set_style_font(image, font_name=BODY_FONT, size=1)
    image.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    image.paragraph_format.space_before = Pt(1)
    image.paragraph_format.space_after = Pt(4)
    image.paragraph_format.keep_together = True

    answer = _paragraph_style(document, "Practice Answer Line")
    _set_style_font(answer, font_name=BODY_FONT, size=10)
    answer.paragraph_format.left_indent = Mm(8)
    answer.paragraph_format.right_indent = Mm(2)
    # 小学生手写一行约需 9mm；space_before 拉开相邻横线（12+10+3=25pt≈8.8mm），
    # 也让第一条横线与题干/题图留出书写起笔空间。
    answer.paragraph_format.space_before = Pt(12)
    answer.paragraph_format.space_after = Pt(3)
    answer.paragraph_format.line_spacing = 1.0
    answer.paragraph_format.keep_together = True


def _apply_page_setup(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(13)
    section.bottom_margin = Mm(13)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)


def _set_paragraph_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "777777")
    borders.append(bottom)


def _append_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run("- ")
    _set_run_font(prefix, font_name=BODY_FONT, size=10, color=GRAY)

    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run._r.extend((begin, instruction, separate, display, end))
    _set_run_font(field_run, font_name=BODY_FONT, size=10, color=GRAY)

    suffix = paragraph.add_run(" -")
    _set_run_font(suffix, font_name=BODY_FONT, size=10, color=GRAY)


def _set_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    _append_page_field(paragraph)


def _set_update_fields(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _emu(mm: float) -> int:
    return int(mm * 36_000)


def _add_page_code_textbox(paragraph, page_code: str, *, shape_id: int) -> None:
    safe_code = escape(page_code)
    box_width = 48 if len(page_code) > 16 else 42
    box_left = 202 - box_width
    anchor = parse_xml(
        f"""
        <wp:anchor
          xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
          xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
          xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
          xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
          distT="0" distB="0" distL="0" distR="0"
          simplePos="0" relativeHeight="{251_660_000 + shape_id}"
          behindDoc="0" locked="1" layoutInCell="0" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="page">
            <wp:posOffset>{_emu(box_left)}</wp:posOffset>
          </wp:positionH>
          <wp:positionV relativeFrom="page"><wp:posOffset>{_emu(7)}</wp:posOffset></wp:positionV>
          <wp:extent cx="{_emu(box_width)}" cy="{_emu(8)}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:wrapNone/>
          <wp:docPr id="{shape_id}" name="Practice page code {shape_id}"/>
          <wp:cNvGraphicFramePr/>
          <a:graphic>
            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
              <wps:wsp>
                <wps:cNvSpPr txBox="1"/>
                <wps:spPr>
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="{_emu(box_width)}" cy="{_emu(8)}"/>
                  </a:xfrm>
                  <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                  <a:noFill/>
                  <a:ln><a:noFill/></a:ln>
                </wps:spPr>
                <wps:txbx>
                  <w:txbxContent>
                    <w:p>
                      <w:pPr><w:jc w:val="center"/></w:pPr>
                      <w:r>
                        <w:rPr>
                          <w:rFonts w:ascii="{CODE_FONT}" w:hAnsi="{CODE_FONT}"
                            w:eastAsia="{CODE_FONT}"/>
                          <w:b/>
                          <w:color w:val="{BLACK}"/>
                          <w:sz w:val="21"/><w:szCs w:val="21"/>
                        </w:rPr>
                        <w:t>{safe_code}</w:t>
                      </w:r>
                    </w:p>
                  </w:txbxContent>
                </wps:txbx>
                <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"
                  horzOverflow="overflow" vert="horz" wrap="square"
                  lIns="36000" tIns="18000" rIns="36000" bIns="18000"
                  numCol="1" spcCol="0" rtlCol="0" anchor="ctr" anchorCtr="1">
                  <a:noAutofit/>
                </wps:bodyPr>
              </wps:wsp>
            </a:graphicData>
          </a:graphic>
        </wp:anchor>
        """
    )
    drawing = OxmlElement("w:drawing")
    drawing.append(anchor)
    run = paragraph.add_run()
    run._r.append(drawing)


def _add_dynamic_page_code_textbox(
    paragraph,
    page_code_prefix: str,
    *,
    shape_id: int,
) -> None:
    safe_prefix = escape(page_code_prefix)
    box_width = 48 if len(page_code_prefix) > 16 else 42
    box_left = 202 - box_width
    anchor = parse_xml(
        f"""
        <wp:anchor
          xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
          xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
          xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
          xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
          distT="0" distB="0" distL="0" distR="0"
          simplePos="0" relativeHeight="{251_660_000 + shape_id}"
          behindDoc="0" locked="1" layoutInCell="0" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="page">
            <wp:posOffset>{_emu(box_left)}</wp:posOffset>
          </wp:positionH>
          <wp:positionV relativeFrom="page"><wp:posOffset>{_emu(7)}</wp:posOffset></wp:positionV>
          <wp:extent cx="{_emu(box_width)}" cy="{_emu(8)}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:wrapNone/>
          <wp:docPr id="{shape_id}" name="Dynamic practice page code {shape_id}"/>
          <wp:cNvGraphicFramePr/>
          <a:graphic>
            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
              <wps:wsp>
                <wps:cNvSpPr txBox="1"/>
                <wps:spPr>
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="{_emu(box_width)}" cy="{_emu(8)}"/>
                  </a:xfrm>
                  <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                  <a:noFill/>
                  <a:ln><a:noFill/></a:ln>
                </wps:spPr>
                <wps:txbx>
                  <w:txbxContent>
                    <w:p>
                      <w:pPr><w:jc w:val="center"/></w:pPr>
                      <w:r>
                        <w:rPr>
                          <w:rFonts w:ascii="{CODE_FONT}" w:hAnsi="{CODE_FONT}"
                            w:eastAsia="{CODE_FONT}"/>
                          <w:b/>
                          <w:color w:val="{BLACK}"/>
                          <w:sz w:val="21"/><w:szCs w:val="21"/>
                        </w:rPr>
                        <w:t>{safe_prefix}</w:t>
                      </w:r>
                      <w:fldSimple w:instr=" PAGE ">
                        <w:r>
                          <w:rPr>
                            <w:rFonts w:ascii="{CODE_FONT}" w:hAnsi="{CODE_FONT}"
                              w:eastAsia="{CODE_FONT}"/>
                            <w:b/>
                            <w:color w:val="{BLACK}"/>
                            <w:sz w:val="21"/><w:szCs w:val="21"/>
                          </w:rPr>
                          <w:t>1</w:t>
                        </w:r>
                      </w:fldSimple>
                    </w:p>
                  </w:txbxContent>
                </wps:txbx>
                <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"
                  horzOverflow="overflow" vert="horz" wrap="square"
                  lIns="36000" tIns="18000" rIns="36000" bIns="18000"
                  numCol="1" spcCol="0" rtlCol="0" anchor="ctr" anchorCtr="1">
                  <a:noAutofit/>
                </wps:bodyPr>
              </wps:wsp>
            </a:graphicData>
          </a:graphic>
        </wp:anchor>
        """
    )
    drawing = OxmlElement("w:drawing")
    drawing.append(anchor)
    run = paragraph.add_run()
    run._r.append(drawing)


def _set_dynamic_page_code_header(section, page_code_prefix: str) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    _add_dynamic_page_code_textbox(paragraph, page_code_prefix, shape_id=2001)
    spacer = paragraph.add_run("\u00a0")
    _set_run_font(spacer, font_name=BODY_FONT, size=1)


def _format_title(batch_date: datetime, *, personal: bool = False) -> str:
    qualifier = "个人" if personal else ""
    return f"{batch_date.month}月{batch_date.day}日{qualifier}练习纸"


def _add_first_page_heading(
    document: Document,
    manifest: PracticeManifest,
    *,
    include_legacy_page_code: bool,
) -> None:
    title = document.add_paragraph(style="Practice Title")
    if include_legacy_page_code:
        _add_page_code_textbox(title, manifest.pages[0].page_code, shape_id=1001)
    title_run = title.add_run(
        _format_title(manifest.batch_date, personal=manifest.student is not None)
    )
    _set_run_font(title_run, font_name=TITLE_FONT, size=16, bold=True)

    student = document.add_paragraph(style="Practice Student Info")
    if manifest.student is None:
        identity = "姓名：________________    学号：________"
    else:
        identity = f"姓名：{manifest.student.name}    学号：{manifest.student.student_number}"
    student_run = student.add_run(identity)
    _set_run_font(student_run, font_name=BODY_FONT, size=12)


def _add_continuation_marker(document: Document, page: PracticePage) -> None:
    marker = document.add_paragraph()
    marker.paragraph_format.space_before = Pt(0)
    marker.paragraph_format.space_after = Pt(5)
    marker.paragraph_format.line_spacing = Pt(1)
    _add_page_code_textbox(marker, page.page_code, shape_id=1000 + page.page_number)
    spacer = marker.add_run("\u00a0")
    _set_run_font(spacer, font_name=BODY_FONT, size=1)


def _resize_picture(shape) -> None:
    scale = min(
        1.0,
        float(MAX_IMAGE_WIDTH) / float(shape.width),
        float(MAX_IMAGE_HEIGHT) / float(shape.height),
    )
    if scale < 1.0:
        shape.width = int(shape.width * scale)
        shape.height = int(shape.height * scale)


def _add_stem_image(document: Document, image: StemImage, *, keep_with_next: bool) -> None:
    paragraph = document.add_paragraph(style="Practice Stem Image")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run()
    shape = run.add_picture(str(image.path))
    _resize_picture(shape)
    shape._inline.docPr.set("descr", image.description)
    shape._inline.docPr.set("title", image.description[:120])


def _add_answer_lines(document: Document, count: int) -> None:
    for line_index in range(count):
        answer = document.add_paragraph(style="Practice Answer Line")
        answer.paragraph_format.keep_with_next = line_index < count - 1
        run = answer.add_run("\u00a0")
        _set_run_font(run, font_name=BODY_FONT, size=10)
        _set_paragraph_bottom_border(answer)


def _add_item(document: Document, item: PracticeItem) -> None:
    paragraph = document.add_paragraph(style="Practice Question")
    paragraph.paragraph_format.keep_with_next = item.stem_image is not None or item.answer_lines > 0
    number = paragraph.add_run(f"{item.display_number}.  ")
    _set_run_font(number, font_name=BODY_FONT, size=12, bold=True)
    question = paragraph.add_run(item.question)
    _set_run_font(question, font_name=BODY_FONT, size=12)

    if item.stem_image is not None:
        _add_stem_image(document, item.stem_image, keep_with_next=item.answer_lines > 0)
    _add_answer_lines(document, item.answer_lines)


def _validate_batch_code(value: str) -> tuple[str, datetime]:
    batch_code = value.strip()
    match = re.fullmatch(r"(?P<date>\d{8})-(?P<sequence>0[1-9]|[1-9]\d)", batch_code)
    if match is None:
        raise ManifestError("批次码必须使用 YYYYMMDD-NN，例如 20260715-01。")
    try:
        batch_date = datetime.strptime(match.group("date"), "%Y%m%d")
    except ValueError as exc:
        raise ManifestError("批次码中的日期无效。") from exc
    return batch_code, batch_date


def _page_code(
    batch_code: str,
    page_number: int,
    student: PracticeStudent | None = None,
) -> str:
    instance = f"-{student.instance_code}" if student is not None else ""
    return f"{batch_code}{instance}-P{page_number}"


def _page_code_prefix(
    batch_code: str,
    student: PracticeStudent | None = None,
) -> str:
    instance = f"-{student.instance_code}" if student is not None else ""
    return f"{batch_code}{instance}-P"


def _required_text(value: object, field_name: str) -> str:
    if not isinstance(value, str) or not value.strip():
        raise ManifestError(f"{field_name} 必须是非空文本。")
    return value.strip()


def _parse_student(raw: object) -> PracticeStudent:
    if not isinstance(raw, Mapping) or set(raw) != {
        "name",
        "studentNumber",
        "instanceCode",
    }:
        raise ManifestError("student 必须包含 name、studentNumber 和 instanceCode。")
    name = _required_text(raw["name"], "student.name")
    student_number = _required_text(raw["studentNumber"], "student.studentNumber")
    instance_code = _required_text(raw["instanceCode"], "student.instanceCode")
    if len(name) > 40 or len(student_number) > 20:
        raise ManifestError("学生姓名或学号过长。")
    if any(ord(character) < 32 for character in name + student_number):
        raise ManifestError("学生姓名和学号不能包含控制字符。")
    if re.fullmatch(r"S\d{3}", instance_code) is None:
        raise ManifestError("student.instanceCode 必须使用 S 加三位数字，例如 S001。")
    return PracticeStudent(name, student_number, instance_code)


def _parse_source(raw: object) -> PracticeSource:
    if not isinstance(raw, Mapping):
        raise ManifestError("source 必须是对象。")
    source_type = raw.get("type")
    if source_type == "original" and set(raw) == {"type", "questionId"}:
        return PracticeSource(
            source_type="original",
            question_id=_required_text(raw["questionId"], "source.questionId"),
            variant_id=None,
        )
    if source_type == "variant" and set(raw) == {"type", "questionId", "variantId"}:
        return PracticeSource(
            source_type="variant",
            question_id=_required_text(raw["questionId"], "source.questionId"),
            variant_id=_required_text(raw["variantId"], "source.variantId"),
        )
    raise ManifestError("source 必须是完整的 original 或 variant 来源。")


def _parse_stem_image(raw: object, manifest_dir: Path) -> StemImage | None:
    if raw is None:
        return None
    if not isinstance(raw, Mapping) or set(raw) != {"path", "description"}:
        raise ManifestError("stemImage 必须包含 path 和 description。")
    raw_path = _required_text(raw["path"], "stemImage.path")
    description = _required_text(raw["description"], "stemImage.description")
    candidate = Path(raw_path).expanduser()
    if not candidate.is_absolute():
        candidate = manifest_dir / candidate
    try:
        resolved = candidate.resolve(strict=True)
    except OSError as exc:
        raise ManifestError("题干图片不可读取。") from exc
    if not resolved.is_file() or resolved.suffix.lower() not in SUPPORTED_IMAGE_SUFFIXES:
        raise ManifestError("题干图片必须是可读取的 JPG、JPEG 或 PNG 文件。")
    return StemImage(path=resolved, description=description)


def _parse_items(
    raw_items: object,
    *,
    manifest_dir: Path,
    first_item_number: int,
) -> tuple[tuple[PracticeItem, ...], int]:
    if not isinstance(raw_items, list) or not raw_items:
        raise ManifestError("items 必须是非空数组。")

    items: list[PracticeItem] = []
    expected_item_number = first_item_number
    for item_raw in raw_items:
        if not isinstance(item_raw, dict):
            raise ManifestError("练习题结构不正确。")
        required_item = {"itemCode", "question", "answerLines", "source"}
        if not required_item.issubset(item_raw) or set(item_raw) - (required_item | {"stemImage"}):
            raise ManifestError("练习题结构不正确。")

        expected_code = f"R{expected_item_number:02d}"
        item_code = _required_text(item_raw["itemCode"], "itemCode")
        question = _required_text(item_raw["question"], "question")
        answer_lines = item_raw["answerLines"]
        if item_code != expected_code:
            raise ManifestError("itemCode 必须从 R01 连续排列。")
        if (
            isinstance(answer_lines, bool)
            or not isinstance(answer_lines, int)
            or not 0 <= answer_lines <= 8
        ):
            raise ManifestError("answerLines 必须是 0 到 8 的整数。")

        items.append(
            PracticeItem(
                item_code=item_code,
                display_number=expected_item_number,
                question=question,
                answer_lines=answer_lines,
                source=_parse_source(item_raw["source"]),
                stem_image=_parse_stem_image(item_raw.get("stemImage"), manifest_dir),
            )
        )
        expected_item_number += 1
    return tuple(items), expected_item_number


def _load_manifest(path: Path) -> PracticeManifest:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ManifestError("无法读取练习纸 manifest。") from exc
    if not isinstance(raw, dict):
        raise ManifestError("manifest 顶层必须是对象。")

    raw_version = raw.get("manifestVersion")
    natural_flow = raw_version in {STANDARD_MANIFEST_VERSION, PERSONAL_MANIFEST_VERSION}
    if raw_version in {LEGACY_STANDARD_MANIFEST_VERSION, STANDARD_MANIFEST_VERSION}:
        student = None
    elif raw_version in {LEGACY_PERSONAL_MANIFEST_VERSION, PERSONAL_MANIFEST_VERSION}:
        student = _parse_student(raw.get("student"))
    else:
        raise ManifestError("manifestVersion 必须是受支持的正式练习纸版本。")

    content_field = "items" if natural_flow else "pages"
    required = {"batchCode", "manifestVersion", content_field}
    if student is not None:
        required.add("student")
    if set(raw) != required:
        raise ManifestError("manifest 字段不完整或包含未知字段。")

    batch_code, batch_date = _validate_batch_code(str(raw["batchCode"]))
    version = _required_text(raw["manifestVersion"], "manifestVersion")
    manifest_dir = path.parent.resolve()
    pages: list[PracticePage] = []
    all_items: list[PracticeItem] = []

    if natural_flow:
        items, _ = _parse_items(
            raw["items"],
            manifest_dir=manifest_dir,
            first_item_number=1,
        )
        all_items.extend(items)
    else:
        pages_raw = raw["pages"]
        if not isinstance(pages_raw, list) or not pages_raw:
            raise ManifestError("pages 必须是非空数组。")
        expected_item_number = 1
        for expected_page, page_raw in enumerate(pages_raw, start=1):
            if not isinstance(page_raw, dict) or set(page_raw) != {"pageNumber", "items"}:
                raise ManifestError("页面结构不正确。")
            if page_raw["pageNumber"] != expected_page:
                raise ManifestError("pageNumber 必须从 1 连续排列。")
            page_items, expected_item_number = _parse_items(
                page_raw["items"],
                manifest_dir=manifest_dir,
                first_item_number=expected_item_number,
            )
            all_items.extend(page_items)
            pages.append(
                PracticePage(
                    page_number=expected_page,
                    page_code=_page_code(batch_code, expected_page, student),
                    items=page_items,
                )
            )

    return PracticeManifest(
        batch_code=batch_code,
        batch_date=batch_date,
        version=version,
        student=student,
        items=tuple(all_items),
        pages=tuple(pages),
        natural_flow=natural_flow,
        page_code_prefix=_page_code_prefix(batch_code, student),
    )


def build_document(manifest_path: Path, output_path: Path) -> None:
    manifest = _load_manifest(manifest_path)
    document = Document()
    _configure_styles(document)
    section = document.sections[0]
    _apply_page_setup(section)
    _set_footer(section)
    _set_update_fields(document)
    if manifest.natural_flow:
        _set_dynamic_page_code_header(section, manifest.page_code_prefix)

    title = _format_title(
        manifest.batch_date,
        personal=manifest.student is not None,
    )
    document.core_properties.title = title
    document.core_properties.subject = (
        "小学数学个人精准练习纸" if manifest.student is not None else "小学数学正式练习纸"
    )
    document.core_properties.author = "shi-homework2lark"

    _add_first_page_heading(
        document,
        manifest,
        include_legacy_page_code=not manifest.natural_flow,
    )
    if manifest.natural_flow:
        for item in manifest.items:
            _add_item(document, item)
    else:
        for page_index, page in enumerate(manifest.pages):
            if page_index > 0:
                document.add_page_break()
                _add_continuation_marker(document, page)
            for item in page.items:
                _add_item(document, item)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Build the formal daily practice DOCX")
    parser.add_argument("--manifest", required=True)
    parser.add_argument("--output", required=True)
    return parser


def _force_utf8_stdio() -> None:
    """Windows 控制台/管道默认 GBK；AI 消费的中文 JSON 必须始终按 UTF-8 输出。"""
    for stream in (sys.stdout, sys.stderr):
        reconfigure = getattr(stream, "reconfigure", None)
        if reconfigure is None:
            continue
        try:
            reconfigure(encoding="utf-8")
        except (ValueError, OSError):
            pass


def main(argv: Sequence[str] | None = None) -> int:
    _force_utf8_stdio()
    args = build_parser().parse_args(argv)
    try:
        manifest_path = Path(args.manifest).resolve()
        build_document(manifest_path, Path(args.output).resolve())
        print(json.dumps({"ok": True, "output": str(Path(args.output))}, ensure_ascii=False))
        return 0
    except ManifestError as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
